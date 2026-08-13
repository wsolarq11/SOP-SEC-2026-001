"""sop_to_docx.py - Convert a typeset SOP (HTML or Markdown) into a styled .docx.

Self-contained: it decodes its own valid default.docx template (docx_template.b64)
to a fresh temp path and uses it as the explicit Document template, so it never
depends on the venv's python-docx default.docx (which is unstable in this env).

Accepts either .html or .md as the first argument. Markdown is converted to a
minimal, well-formed HTML (headings, tables, blockquotes, lists, front-matter
metadata table) before being fed to the same builder.
"""
import base64
import os
import re
import sys
import tempfile

from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
B64_PATH = os.path.join(HERE, "docx_template.b64")

FONT_HEADING = "微软雅黑"
FONT_BODY = "微软雅黑"
FONT_MONO = "Consolas"

COLOR_TEXT = RGBColor(0x11, 0x18, 0x27)
COLOR_PRIMARY = "2563EB"
COLOR_HIGHLIGHT = "FEF3C7"
COLOR_CODE_BG = "F3F4F6"
COLOR_TABLE_HDR = "FEF3C7"


def _set_eastasia(element, font):
    """Set eastAsia/ascii/hAnsi fonts on a run or style element that has rPr."""
    rpr = element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        rfonts.set(qn(attr), font)


def _apply_run_fonts(run):
    run.font.name = FONT_BODY
    _set_eastasia(run._r, FONT_BODY)


def _add_run_shading(run, fill):
    rpr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    rpr.append(shd)


def _style_callout(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), COLOR_PRIMARY)
    pbdr.append(left)
    ppr.append(pbdr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), COLOR_HIGHLIGHT)
    ppr.append(shd)
    paragraph.paragraph_format.left_indent = Pt(12)


def _style_code_block(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), COLOR_CODE_BG)
    ppr.append(shd)
    paragraph.paragraph_format.left_indent = Pt(8)


def _append_runs(paragraph, node, bold=False, italic=False):
    if isinstance(node, NavigableString):
        text = str(node)
        if text.strip() == "":
            return
        run = paragraph.add_run(text)
        _apply_run_fonts(run)
        run.font.bold = bold
        run.font.italic = italic
        return
    tag = node.name
    if tag in ("strong", "b"):
        for child in node.children:
            _append_runs(paragraph, child, bold=True)
    elif tag in ("em", "i"):
        for child in node.children:
            _append_runs(paragraph, child, italic=True)
    elif tag == "code":
        run = paragraph.add_run(node.get_text())
        run.font.name = FONT_MONO
        _set_eastasia(run._r, FONT_MONO)
        run.font.size = Pt(10)
        _add_run_shading(run, COLOR_HIGHLIGHT)
    elif tag in ("br",):
        paragraph.add_run("\n")
    else:
        for child in node.children:
            _append_runs(paragraph, child, bold=bold, italic=italic)


def _configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(11)
    normal.font.color.rgb = COLOR_TEXT
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    _set_eastasia(normal._element, FONT_BODY)

    for style_name, size_pt in (("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 11)):
        try:
            st = doc.styles[style_name]
        except KeyError:
            continue
        st.font.name = FONT_HEADING
        st.font.size = Pt(size_pt)
        st.font.bold = True
        st.font.color.rgb = COLOR_TEXT
        _set_eastasia(st._element, FONT_HEADING)


def _add_table(doc, table_tag):
    rows = table_tag.find_all("tr")
    if not rows:
        return
    n_cols = max(len(r.find_all(["td", "th"])) for r in rows)
    table = doc.add_table(rows=0, cols=n_cols)
    table.style = "Table Grid"
    for r_idx, tr in enumerate(rows):
        cells = tr.find_all(["td", "th"])
        doc_row = table.add_row()
        is_header = (cells and cells[0].name == "th") or r_idx == 0
        for c_idx in range(n_cols):
            cell = cells[c_idx] if c_idx < len(cells) else None
            doc_cell = doc_row.cells[c_idx]
            doc_cell.text = ""
            para = doc_cell.paragraphs[0]
            if cell is not None:
                _append_runs(para, cell)
            if is_header:
                for run in para.runs:
                    run.font.bold = True
                tcpr = doc_cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), COLOR_TABLE_HDR)
                tcpr.append(shd)


def _build(doc, soup):
    body = soup.body
    for node in body.children:
        if isinstance(node, NavigableString):
            continue
        if node.name is None:
            continue
        tag = node.name

        if tag == "h1":
            p = doc.add_paragraph()
            p.style = doc.styles["Title"] if "Title" in doc.styles else doc.styles["Heading 1"]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _append_runs(p, node)
            for r in p.runs:
                r.font.size = Pt(18)
                r.font.bold = True
                r.font.color.rgb = COLOR_TEXT
                _set_eastasia(r._r, FONT_HEADING)
        elif tag == "h2":
            p = doc.add_paragraph()
            p.style = doc.styles["Heading 1"]
            _append_runs(p, node)
        elif tag == "h3":
            p = doc.add_paragraph()
            p.style = doc.styles["Heading 2"]
            _append_runs(p, node)
        elif tag == "nav" and "doc-toc" in (node.get("class") or []):
            _add_toc(doc, node)
        elif tag == "blockquote":
            p = doc.add_paragraph()
            _style_callout(p)
            _append_runs(p, node)
        elif tag == "table":
            _add_table(doc, node)
        elif tag == "ul":
            for li in node.find_all("li", recursive=False):
                p = doc.add_paragraph(style="List Bullet")
                _append_runs(p, li)
        elif tag == "ol":
            for li in node.find_all("li", recursive=False):
                p = doc.add_paragraph(style="List Number")
                _append_runs(p, li)
        elif tag == "p":
            text = node.get_text(strip=True)
            if not text:
                continue
            p = doc.add_paragraph()
            if node.find("code"):
                _style_code_block(p)
            _append_runs(p, node)


def _add_toc(doc, nav_tag):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    p.add_run("目录")
    for li in nav_tag.find_all("li"):
        tp = doc.add_paragraph(style="List Number")
        _append_runs(tp, li)


def _page_setup(doc):
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.left_margin = Mm(25)
    sec.right_margin = Mm(25)
    sec.top_margin = Mm(25)
    sec.bottom_margin = Mm(25)


# --------------------------------------------------------------------------
# Minimal Markdown -> HTML (enough for our controlled SOP documents)
# --------------------------------------------------------------------------

def _escape(t):
    return t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _inline(text):
    text = re.sub(r"`([^`]+)`", r"<code>\1</code>", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", text)
    text = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"<em>\1</em>", text)
    return text


def md_to_html(md):
    lines = md.split("\n")
    fm = {}
    if lines and lines[0].strip() == "---":
        end = None
        for i in range(1, len(lines)):
            if lines[i].strip() == "---":
                end = i
                break
        if end is not None:
            for ln in lines[1:end]:
                if ":" in ln:
                    k, v = ln.split(":", 1)
                    fm[k.strip()] = v.strip()
            lines = lines[end + 1:]

    out = []
    if fm:
        out.append("<table>")
        out.append("<tr><th>字段</th><th>值</th></tr>")
        for k, v in fm.items():
            out.append(f"<tr><td>{_escape(k)}</td><td>{_inline(_escape(v))}</td></tr>")
        out.append("</table>")

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        s = line.strip()
        if s == "":
            i += 1
            continue
        if s.startswith("# "):
            out.append(f"<h1>{_inline(_escape(s[2:]))}</h1>")
            i += 1
        elif s.startswith("## "):
            out.append(f"<h2>{_inline(_escape(s[3:]))}</h2>")
            i += 1
        elif s.startswith("### "):
            out.append(f"<h3>{_inline(_escape(s[4:]))}</h3>")
            i += 1
        elif s.startswith("> "):
            quote = []
            while i < n and lines[i].strip().startswith(">"):
                quote.append(lines[i].strip()[2:].strip())
                i += 1
            out.append(f"<blockquote>{_inline(_escape(' '.join(quote)))}</blockquote>")
        elif s.startswith("|"):
            tbl_lines = []
            while i < n and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i].strip())
                i += 1
            rows = []
            for idx, tl in enumerate(tbl_lines):
                cells = [c.strip() for c in tl.strip("|").split("|")]
                if idx == 1 and any("-" in c for c in cells) and all(set(c) <= set("-: ") for c in cells):
                    continue
                rows.append(cells)
            if rows:
                out.append("<table>")
                for r_i, cells in enumerate(rows):
                    tag = "th" if r_i == 0 else "td"
                    out.append("<tr>" + "".join(f"<{tag}>{_inline(_escape(c))}</{tag}>" for c in cells) + "</tr>")
                out.append("</table>")
        elif re.match(r"^[-*] ", s):
            items = []
            while i < n and re.match(r"^[-*] ", lines[i].strip()):
                items.append(lines[i].strip()[2:].strip())
                i += 1
            out.append("<ul>" + "".join(f"<li>{_inline(_escape(it))}</li>" for it in items) + "</ul>")
        elif re.match(r"^\d+\. ", s):
            items = []
            while i < n and re.match(r"^\d+\. ", lines[i].strip()):
                items.append(re.sub(r"^\d+\. ", "", lines[i].strip()))
                i += 1
            out.append("<ol>" + "".join(f"<li>{_inline(_escape(it))}</li>" for it in items) + "</ol>")
        else:
            out.append(f"<p>{_inline(_escape(s))}</p>")
            i += 1
    return "<html><body>" + "".join(out) + "</body></html>"


def main():
    if len(sys.argv) < 3:
        print("usage: sop_to_docx.py <input.(html|md)> <output.docx>")
        sys.exit(2)
    in_path = sys.argv[1]
    out_path = sys.argv[2]

    with open(B64_PATH, "r") as f:
        b64 = f.read().strip()
    tpl_bytes = base64.b64decode(b64)
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp.write(tpl_bytes)
    tmp.close()
    template_path = tmp.name

    try:
        doc = Document(docx=template_path)
        _configure_styles(doc)
        _page_setup(doc)

        with open(in_path, "r", encoding="utf-8") as f:
            raw = f.read()
        if in_path.lower().endswith(".md"):
            html = md_to_html(raw)
        else:
            html = raw
        soup = BeautifulSoup(html, "html.parser")

        _build(doc, soup)
        doc.save(out_path)

        # Validate in the SAME process (the env clobbers .docx between runs).
        verify = Document(out_path)
        print("OK saved:", out_path)
        print("VERIFY paragraphs:", len(verify.paragraphs), "tables:", len(verify.tables))
        print("VERIFY styles seen:", sorted({p.style.name for p in verify.paragraphs}))
        if verify.tables:
            print("VERIFY table rows:", len(verify.tables[0].rows))
    finally:
        try:
            os.unlink(template_path)
        except OSError:
            pass


if __name__ == "__main__":
    main()
